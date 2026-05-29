from __future__ import annotations

import asyncio
import csv
import json
import re
import time
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List

import requests
from fastapi import HTTPException
from sqlalchemy.orm import Session

from ..core.config import settings
from ..models.database import KnowledgeDocument
from ..rag.embedding import BGEEmbedder
from ..utils.logger import logger
from ..vector_store.milvus_client import MilvusClient

TEXT_EXTENSIONS = {".txt", ".md", ".json", ".csv"}
DOCX_EXTENSIONS = {".docx"}
PDF_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".gif", ".jp2"}
OFFICE_MINERU_EXTENSIONS = {".doc", ".ppt", ".pptx", ".xls", ".xlsx"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | DOCX_EXTENSIONS | PDF_EXTENSIONS | IMAGE_EXTENSIONS | OFFICE_MINERU_EXTENSIONS


@dataclass
class ParsedDocument:
    content: str
    parser: str
    category: str
    use_mineru: bool
    reason: str


class MinerUClient:
    def __init__(self) -> None:
        self.base_url = settings.MINERU_API_URL.rstrip("/")
        self.timeout = settings.MINERU_REQUEST_TIMEOUT_SECONDS
        self.poll_interval = max(1, settings.MINERU_POLL_INTERVAL_SECONDS)

    def parse_file(self, file_path: Path) -> str:
        payload = {
            "file_name": file_path.name,
            "language": "ch",
            "enable_table": True,
            "enable_formula": True,
            "is_ocr": True,
        }

        create_url = f"{self.base_url}/api/v1/agent/parse/file"
        logger.info("MinerU task create: file=%s url=%s", file_path, create_url)
        create_response = requests.post(create_url, json=payload, timeout=self.timeout)
        create_response.raise_for_status()
        create_data = create_response.json()

        task_data = create_data.get("data") if isinstance(create_data, dict) else {}
        task_id = task_data.get("task_id") or task_data.get("id")
        upload_url = task_data.get("file_url") or task_data.get("upload_url")
        if not task_id or not upload_url:
            raise RuntimeError(f"MinerU returned invalid task payload: {create_data}")

        with file_path.open("rb") as handle:
            upload_response = requests.put(upload_url, data=handle, timeout=self.timeout)
        upload_response.raise_for_status()
        logger.info("MinerU file uploaded: file=%s task_id=%s", file_path, task_id)

        return self._poll_result(task_id)

    def _poll_result(self, task_id: str) -> str:
        query_url = f"{self.base_url}/api/v1/agent/parse/{task_id}"
        deadline = time.time() + self.timeout

        while time.time() < deadline:
            response = requests.get(query_url, timeout=self.timeout)
            response.raise_for_status()
            payload = response.json()
            data = payload.get("data") if isinstance(payload, dict) and isinstance(payload.get("data"), dict) else payload
            state = str(
                data.get("state")
                or data.get("status")
                or payload.get("state")
                or payload.get("status")
                or ""
            ).lower()

            if state == "done":
                markdown_url = data.get("markdown_url") or data.get("md_url") or data.get("full_md_url")
                if markdown_url:
                    markdown_response = requests.get(markdown_url, timeout=self.timeout)
                    markdown_response.raise_for_status()
                    return markdown_response.text.strip()
                return str(data.get("markdown") or data.get("content") or "").strip()

            if state in {"failed", "error"}:
                raise RuntimeError(f"MinerU parsing failed: {payload}")

            time.sleep(self.poll_interval)

        raise TimeoutError(f"MinerU parsing timed out for task_id={task_id}")


class DocumentComplexityClassifier:
    TABLE_LINE_PATTERN = re.compile(r"(\|.+\|)|(\t)|([,，;；]\s*[^,，;；\n]+[,，;；]\s*[^,，;；\n]+)")
    FORMULA_PATTERN = re.compile(
        r"(∑|∫|√|∞|≤|≥|≠|≈|±|÷|×|∂|∇|λ|μ|σ|π|θ|β|α|γ|Δ|Ω|Φ|Ψ|≡|→|←|⇌|:=|\\frac|\\sum|\\int|\b[a-zA-Z]\^\d+\b)"
    )

    def classify(self, *, file_path: Path, text: str, parser: str) -> tuple[str, bool, str]:
        suffix = file_path.suffix.lower()
        normalized = self._normalize_text(text)

        if suffix in IMAGE_EXTENSIONS:
            return "image", True, "image file requires OCR/layout parsing"

        if suffix in OFFICE_MINERU_EXTENSIONS:
            return "complex-office", True, "office document routed to MinerU"

        if suffix in TEXT_EXTENSIONS:
            return "plain-text", False, "plain text format"

        if suffix in DOCX_EXTENSIONS:
            table_score = self._table_score(normalized)
            formula_score = self._formula_score(normalized)
            if table_score >= settings.DOC_TABLE_TRIGGER_SCORE:
                return "table-heavy", True, f"table score {table_score} >= {settings.DOC_TABLE_TRIGGER_SCORE}"
            if formula_score >= settings.DOC_FORMULA_TRIGGER_SCORE:
                return "formula-heavy", True, f"formula score {formula_score} >= {settings.DOC_FORMULA_TRIGGER_SCORE}"
            return "docx-text", False, "docx parsed locally"

        if suffix in PDF_EXTENSIONS:
            if not normalized:
                return "scan-pdf", True, "local PDF extraction empty"
            if len(normalized) < settings.MINERU_PDF_FALLBACK_MIN_CHARS:
                return "short-pdf", True, "local PDF text too short"
            if "\ufffd" in normalized or normalized.count("?") >= 6:
                return "garbled-pdf", True, "local PDF text looks garbled"

            table_score = self._table_score(normalized)
            formula_score = self._formula_score(normalized)
            if table_score >= settings.PDF_TABLE_TRIGGER_SCORE:
                return "table-heavy", True, f"table score {table_score} >= {settings.PDF_TABLE_TRIGGER_SCORE}"
            if formula_score >= settings.PDF_FORMULA_TRIGGER_SCORE:
                return "formula-heavy", True, f"formula score {formula_score} >= {settings.PDF_FORMULA_TRIGGER_SCORE}"
            return "pdf-text", False, f"local parser={parser}"

        return "fallback", False, "default local flow"

    def _table_score(self, text: str) -> int:
        if not text:
            return 0
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        score = 0
        for line in lines:
            if self.TABLE_LINE_PATTERN.search(line):
                score += 1
            compact_cells = re.split(r"[,\t|，；;]", line)
            meaningful_cells = [cell.strip() for cell in compact_cells if cell.strip()]
            if len(meaningful_cells) >= 4:
                score += 1
        return score

    def _formula_score(self, text: str) -> int:
        if not text:
            return 0
        matches = self.FORMULA_PATTERN.findall(text)
        return len(matches)

    def _normalize_text(self, text: str) -> str:
        normalized = re.sub(r"\r\n?", "\n", text or "")
        normalized = re.sub(r"[ \t]+", " ", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()


class DocumentParser:
    def parse_local(self, file_path: Path) -> tuple[str, str]:
        suffix = file_path.suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {suffix}")

        if suffix in TEXT_EXTENSIONS:
            return self._read_text(file_path), "local-text"
        if suffix in DOCX_EXTENSIONS:
            return self._read_docx(file_path), "local-docx"
        if suffix in PDF_EXTENSIONS:
            return self._read_pdf_local(file_path), "local-pdf"
        if suffix in IMAGE_EXTENSIONS:
            return "", "image-pending-ocr"
        if suffix in OFFICE_MINERU_EXTENSIONS:
            return "", "office-pending-mineru"
        raise ValueError(f"Unsupported file format: {suffix}")

    def _read_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()
        if suffix == ".json":
            with file_path.open("r", encoding="utf-8") as handle:
                return json.dumps(json.load(handle), ensure_ascii=False, indent=2)
        if suffix == ".csv":
            rows: List[str] = []
            with file_path.open("r", encoding="utf-8-sig", newline="") as handle:
                reader = csv.reader(handle)
                for row in reader:
                    joined = ", ".join(cell.strip() for cell in row if cell and cell.strip())
                    if joined:
                        rows.append(joined)
            return "\n".join(rows)

        for encoding in ("utf-8", "utf-8-sig", "gbk", "gb2312"):
            try:
                return file_path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        return file_path.read_text(encoding="utf-8", errors="ignore")

    def _read_docx(self, file_path: Path) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise ValueError("python-docx is required for DOCX files") from exc

        document = Document(str(file_path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n".join(paragraphs)

    def _read_pdf_local(self, file_path: Path) -> str:
        try:
            import fitz

            document = fitz.open(str(file_path))
            try:
                pages = [page.get_text("text") or "" for page in document]
            finally:
                document.close()
            text = self._normalize_text("\n\n".join(pages))
            if text:
                return text
        except Exception as exc:
            logger.warning("PyMuPDF parse failed for %s: %s", file_path, exc)

        try:
            from pypdf import PdfReader
        except ImportError:
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                return ""

        try:
            reader = PdfReader(str(file_path))
            pages = [page.extract_text() or "" for page in reader.pages]
            return self._normalize_text("\n\n".join(pages))
        except Exception as exc:
            logger.warning("PDF text extraction failed for %s: %s", file_path, exc)
            return ""

    def _normalize_text(self, text: str) -> str:
        normalized = re.sub(r"\r\n?", "\n", text or "")
        normalized = re.sub(r"[ \t]+", " ", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()


class DocumentIngestionService:
    def __init__(self, db: Session):
        self.db = db
        self.parser = DocumentParser()
        self.classifier = DocumentComplexityClassifier()
        self.mineru = MinerUClient()
        self.embedder = BGEEmbedder()
        self.milvus = MilvusClient()

    async def ingest_uploaded_document(
        self,
        *,
        title: str,
        doc_type: str,
        legal_field: str,
        source: str,
        file_path: Path,
        original_filename: str | None = None,
    ) -> dict:
        doc = KnowledgeDocument(
            title=title,
            doc_type=doc_type,
            legal_field=legal_field,
            source=source,
            file_path=str(file_path),
            status="processing",
            metadata_json={
                "original_filename": original_filename or file_path.name,
                "extension": file_path.suffix.lower(),
            },
        )
        self.db.add(doc)
        self.db.commit()
        self.db.refresh(doc)

        try:
            parsed = await asyncio.to_thread(self._parse_and_route, file_path)
            chunks = self._split_text(parsed.content)
            if not chunks:
                raise ValueError("Document content is empty after parsing")

            logger.info("Knowledge split completed: file_path=%s chunks=%s", file_path, len(chunks))
            embed_result = await asyncio.to_thread(
                self.embedder.encode,
                chunks,
                batch_size=min(8, len(chunks)),
                return_sparse=False,
            )
            dense_vectors = [vec.tolist() for vec in embed_result["dense_vecs"]]
            logger.info("Embedding completed: file_path=%s chunks=%s", file_path, len(chunks))
            insert_ids = await asyncio.to_thread(
                self.milvus.insert,
                dense_vectors=dense_vectors,
                contents=chunks,
                sources=[source] * len(chunks),
                article_numbers=[""] * len(chunks),
                legal_fields=[legal_field] * len(chunks),
                knowledge_types=[doc_type] * len(chunks),
                document_titles=[title] * len(chunks),
            )
            logger.info("Milvus insert completed: file_path=%s inserted=%s", file_path, len(insert_ids))

            doc.chunk_count = len(chunks)
            doc.status = "completed"
            doc.metadata_json = {
                **(doc.metadata_json or {}),
                "parser": parsed.parser,
                "category": parsed.category,
                "use_mineru": parsed.use_mineru,
                "route_reason": parsed.reason,
                "chunk_count": len(chunks),
                "file_size": file_path.stat().st_size if file_path.exists() else 0,
            }
            self.db.commit()
            self.db.refresh(doc)

            return {
                "id": doc.id,
                "status": doc.status,
                "chunk_count": doc.chunk_count,
                "parser": parsed.parser,
                "category": parsed.category,
                "use_mineru": parsed.use_mineru,
                "route_reason": parsed.reason,
                "inserted": len(insert_ids),
                "message": "Knowledge document ingested successfully",
            }
        except Exception as exc:
            doc.status = "failed"
            doc.metadata_json = {
                **(doc.metadata_json or {}),
                "error": str(exc),
            }
            self.db.commit()
            logger.error("Knowledge ingest failed for %s: %s", file_path, exc)
            raise HTTPException(status_code=500, detail=f"Knowledge ingest failed: {exc}") from exc

    def _parse_and_route(self, file_path: Path) -> ParsedDocument:
        local_text, local_parser = self.parser.parse_local(file_path)
        category, use_mineru, reason = self.classifier.classify(
            file_path=file_path,
            text=local_text,
            parser=local_parser,
        )
        normalized_local = self._normalize_text(local_text)

        if self._should_skip_mineru_for_large_pdf(file_path, normalized_local):
            return ParsedDocument(
                content=normalized_local,
                parser=f"{local_parser}-large-file-fallback",
                category=f"{category}-large-file-fallback",
                use_mineru=False,
                reason=(
                    f"{reason}; skipped MinerU because file size exceeds "
                    f"{settings.MINERU_LIGHTWEIGHT_MAX_FILE_MB}MB lightweight API limit"
                ),
            )

        if use_mineru:
            try:
                mineru_text = self.mineru.parse_file(file_path)
                normalized = self._normalize_text(mineru_text)
                if not normalized:
                    raise ValueError(f"MinerU returned empty content for {file_path.name}")
                return ParsedDocument(
                    content=normalized,
                    parser="mineru",
                    category=category,
                    use_mineru=True,
                    reason=reason,
                )
            except Exception as exc:
                normalized_local = self._normalize_text(local_text)
                if normalized_local:
                    logger.warning(
                        "MinerU parsing failed for %s, falling back to %s: %s",
                        file_path,
                        local_parser,
                        exc,
                    )
                    return ParsedDocument(
                        content=normalized_local,
                        parser=f"{local_parser}-fallback",
                        category=f"{category}-fallback",
                        use_mineru=False,
                        reason=f"{reason}; MinerU fallback to local parser because: {exc}",
                    )
                raise

        if not normalized_local:
            raise ValueError(f"Local parser produced empty content for {file_path.name}")
        return ParsedDocument(
            content=normalized_local,
            parser=local_parser,
            category=category,
            use_mineru=False,
            reason=reason,
        )

    def _should_skip_mineru_for_large_pdf(self, file_path: Path, normalized_local: str) -> bool:
        if file_path.suffix.lower() not in PDF_EXTENSIONS:
            return False
        if not normalized_local:
            return False
        if not file_path.exists():
            return False

        max_bytes = max(settings.MINERU_LIGHTWEIGHT_MAX_FILE_MB, 1) * 1024 * 1024
        return file_path.stat().st_size > max_bytes

    def _split_text(self, text: str) -> List[str]:
        normalized = self._normalize_text(text)
        if not normalized:
            return []

        max_sentence_length = max(settings.CHUNK_SIZE, 200)
        paragraphs = [part.strip() for part in re.split(r"\n{2,}", normalized) if part.strip()]
        chunks: List[str] = []

        for paragraph in paragraphs:
            sentences = self._split_paragraph_into_sentences(paragraph)
            for sentence in sentences:
                if len(sentence) <= max_sentence_length:
                    chunks.append(sentence)
                    continue
                chunks.extend(self._split_long_sentence(sentence, max_sentence_length))

        return self._deduplicate_preserve_order(chunks)

    def _split_paragraph_into_sentences(self, paragraph: str) -> List[str]:
        compact_paragraph = re.sub(r"\s*\n\s*", " ", paragraph).strip()
        if not compact_paragraph:
            return []

        sentence_pattern = re.compile(
            r'[^\u3002\uff01\uff1f!?;\uff1b]+(?:[\u3002\uff01\uff1f!?;\uff1b]+(?:[""\'\'\u201d\u2019\u300d\u300f\u3011\uff09)]*)?)?|[\u3002\uff01\uff1f!?;\uff1b]+'
        )
        sentences = [match.group(0).strip() for match in sentence_pattern.finditer(compact_paragraph)]
        return [sentence for sentence in sentences if sentence]

    def _split_long_sentence(self, sentence: str, max_sentence_length: int) -> List[str]:
        sentence = sentence.strip()
        if not sentence:
            return []
        if len(sentence) <= max_sentence_length:
            return [sentence]

        parts = [part.strip() for part in re.split(r"(?<=[,\u3001:\uff1a\uff0c])", sentence) if part.strip()]
        if len(parts) > 1:
            chunks: List[str] = []
            current = ""
            for part in parts:
                if not current:
                    current = part
                    continue
                if len(current) + len(part) <= max_sentence_length:
                    current += part
                    continue
                chunks.extend(self._split_long_sentence(current, max_sentence_length))
                current = part
            if current:
                chunks.extend(self._split_long_sentence(current, max_sentence_length))
            return chunks

        chunks: List[str] = []
        start = 0
        while start < len(sentence):
            chunk = sentence[start:start + max_sentence_length].strip()
            if chunk:
                chunks.append(chunk)
            start += max_sentence_length
        return chunks

    def _deduplicate_preserve_order(self, items: Iterable[str]) -> List[str]:
        seen = set()
        result = []
        for item in items:
            key = item.strip()
            if not key or key in seen:
                continue
            seen.add(key)
            result.append(key)
        return result

    def _normalize_text(self, text: str) -> str:
        normalized = re.sub(r"\r\n?", "\n", text or "")
        normalized = re.sub(r"[ \t]+", " ", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        return normalized.strip()
